from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Smart_Power_Teknik_Proje_Rehberi.docx"

NAVY = "17365D"
BLUE = "3B73D9"
LIME = "B9F65A"
PALE_BLUE = "EAF2FB"
PALE_GREEN = "EEF8ED"
PALE_YELLOW = "FFF4D6"
INK = "17243A"
MUTED = "52647A"
LINE = "C9D7E8"
WHITE = "FFFFFF"
RED = "A63A45"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = LINE, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def keep_with_next(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    p_pr.append(keep)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_run(paragraph, text: str, *, bold=False, color=INK, size=8.7, italic=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def style_paragraph(paragraph, *, space_before=0, space_after=3, line_spacing=1.03):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    fmt.line_spacing = line_spacing


def add_bullet(doc, text: str, *, level=0, bold_prefix: str | None = None, size=8.6, color=INK):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    style_paragraph(p, space_after=2.2)
    if bold_prefix and text.startswith(bold_prefix):
        add_run(p, bold_prefix, bold=True, color=color, size=size)
        add_run(p, text[len(bold_prefix):], color=color, size=size)
    else:
        add_run(p, text, color=color, size=size)
    return p


def add_numbered(doc, text: str, *, size=8.6):
    p = doc.add_paragraph(style="List Number")
    style_paragraph(p, space_after=2.2)
    add_run(p, text, size=size)
    return p


def add_page_header(doc, section_label: str, title: str, subtitle: str | None = None):
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.columns[0].width = Inches(5.65)
    t.columns[1].width = Inches(1.35)
    for c in t.rows[0].cells:
        set_cell_shading(c, NAVY)
        set_cell_border(c, NAVY, "0")
        set_cell_margins(c, top=95, bottom=95)
    p = t.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, title, bold=True, color=WHITE, size=16)
    q = t.cell(0, 1).paragraphs[0]
    q.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(q, section_label.upper(), bold=True, color=LIME, size=8)
    if subtitle:
        p = doc.add_paragraph()
        style_paragraph(p, space_before=3, space_after=7)
        add_run(p, subtitle, color=MUTED, size=8.6)
    else:
        spacer = doc.add_paragraph()
        style_paragraph(spacer, space_after=2)


def add_section_heading(doc, text: str, *, color=NAVY, size=11.4):
    p = doc.add_paragraph()
    style_paragraph(p, space_before=5, space_after=3)
    keep_with_next(p)
    add_run(p, text, bold=True, color=color, size=size)
    return p


def add_card_grid(doc, cards, *, columns=4, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for idx, card in enumerate(cards):
        cell = table.cell(0, idx)
        set_cell_shading(cell, card.get("fill", fill))
        set_cell_border(cell, card.get("border", LINE), "7")
        set_cell_margins(cell, top=110, start=95, bottom=110, end=95)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(p, space_after=1)
        add_run(p, card["value"], bold=True, color=card.get("color", BLUE), size=15.5)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(p, space_after=0)
        add_run(p, card["label"], bold=True, color=INK, size=7.4)
    return table


def add_callout(doc, title: str, text: str, *, fill=PALE_YELLOW, accent=BLUE):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(0.12)
    table.columns[1].width = Inches(6.88)
    c0, c1 = table.rows[0].cells
    set_cell_shading(c0, accent)
    set_cell_border(c0, accent, "0")
    set_cell_shading(c1, fill)
    set_cell_border(c1, fill, "0")
    set_cell_margins(c1, top=100, start=125, bottom=100, end=125)
    p = c1.paragraphs[0]
    style_paragraph(p, space_after=1)
    add_run(p, title, bold=True, color=NAVY, size=9.2)
    p = c1.add_paragraph()
    style_paragraph(p, space_after=0)
    add_run(p, text, color=INK, size=8.3)
    return table


def add_pipeline(doc):
    labels = [
        ("01", "API ingestion", "4 kaynak · ham CSV"),
        ("02", "Hourly transform", "UTC · ölçü · özellik"),
        ("03", "MySQL + SQL", "şema · view · QA"),
        ("04", "Analysis / ML", "hipotez · hold-out"),
        ("05", "Product", "Streamlit · Tableau"),
    ]
    table = doc.add_table(rows=1, cols=len(labels))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, (num, name, detail) in enumerate(labels):
        cell = table.cell(0, i)
        set_cell_shading(cell, NAVY if i in (0, 2, 4) else BLUE)
        set_cell_border(cell, WHITE, "12")
        set_cell_margins(cell, top=95, start=65, bottom=95, end=65)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, num, bold=True, color=LIME, size=7.8)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(p, space_after=1)
        add_run(p, name, bold=True, color=WHITE, size=7.6)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, detail, color=WHITE, size=6.5)
    return table


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(p, space_before=0, space_after=0)
    add_run(p, "SMART POWER · TEKNİK PROJE REHBERİ     |     ", color=MUTED, size=7)
    run = p.add_run()
    run.font.name = "Aptos"
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    field_code = OxmlElement("w:instrText")
    field_code.set(qn("xml:space"), "preserve")
    field_code.text = " PAGE "
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.extend([field_begin, field_code, field_end])
    add_run(p, "/5", color=MUTED, size=7)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def table_text(cell, text, *, bold=False, color=INK, size=7.6, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    p.alignment = align
    style_paragraph(p, space_after=0)
    add_run(p, text, bold=bold, color=color, size=size)


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(0.48)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.2)
    add_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(8.7)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.03

    # PAGE 1 — memo masthead
    mast = doc.add_table(rows=1, cols=2)
    mast.alignment = WD_TABLE_ALIGNMENT.CENTER
    mast.autofit = False
    mast.columns[0].width = Inches(4.8)
    mast.columns[1].width = Inches(2.2)
    for c in mast.rows[0].cells:
        set_cell_shading(c, NAVY)
        set_cell_border(c, NAVY, "0")
        set_cell_margins(c, top=110, bottom=110)
    p = mast.cell(0, 0).paragraphs[0]
    add_run(p, "SMART POWER", bold=True, color=WHITE, size=13.5)
    p = mast.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p, "TEKNİK MÜLAKAT DOSYASI", bold=True, color=LIME, size=7.6)

    p = doc.add_paragraph()
    style_paragraph(p, space_before=9, space_after=2)
    add_run(p, "Saatlik elektrik verisinden", bold=True, color=NAVY, size=22)
    p = doc.add_paragraph()
    style_paragraph(p, space_after=5)
    add_run(p, "tarihsel karar destek ürününe", bold=True, color=BLUE, size=22)
    p = doc.add_paragraph()
    style_paragraph(p, space_after=9)
    add_run(p, "Teknik ekip için; veri toplama, dönüşüm, SQL, analiz, model ve ürün akışının 5 sayfalık özeti.", color=MUTED, size=9.4)

    add_callout(
        doc,
        "Araştırma sorusu",
        "Hangi dört saatlik zaman aralığı, Hollanda’da esnek elektrik kullanımını hem daha düşük fiyatlı hem de daha düşük karbon yoğunluklu saatlere taşıyor?",
        fill=PALE_BLUE,
        accent=BLUE,
    )

    add_section_heading(doc, "60 saniyelik teknik özet")
    p = doc.add_paragraph()
    style_paragraph(p, space_after=6)
    add_run(
        p,
        "Dört harici kaynaktan fiyat, üretim karması, hava koşulları ve yaşam döngüsü karbon yoğunluğu verisini Python ile aldım. Tüm zamanları UTC’ye çevirdim; 15 dakikalık karbon verisini yalnızca dört tam gözlem bulunan saatler için ortalamaya aldım ve kaynakları saatlik timestamp üzerinden birleştirdim. Temiz tabloyu MySQL’e idempotent bir şema ile yükledim; SQL view ve kalite sorguları oluşturdum. Ardından tarihsel düşük fiyat/düşük karbon kesişimini, eşleşmiş gün karşılaştırmasını ve kronolojik hold-out kullanan bir sonraki saat fiyat modelini değerlendirdim. Sonuçları Streamlit karar destek prototipi ve Tableau dashboard’u ile ürünleştirdim.",
        size=8.8,
    )

    add_card_grid(
        doc,
        [
            {"value": "1,170", "label": "doğrulanmış saat"},
            {"value": "12–16", "label": "önerilen tarihsel pencere"},
            {"value": "164", "label": "ucuz ve düşük karbonlu saat"},
            {"value": "4", "label": "API kaynağı"},
        ],
    )

    add_section_heading(doc, "Ana çıktı ve kapsam")
    add_bullet(doc, "Örneklem: 28 Mayıs–3 Ağustos 2026; Amsterdam sunumu için yerel saate çevrilen, analizde UTC ile hizalanan 1.170 saat.", bold_prefix="Örneklem:")
    add_bullet(doc, "Tarihsel bulgu: 12:00–16:00 penceresi fiyat–karbon dengesinde en iyi dört saatlik blok olarak çıktı.", bold_prefix="Tarihsel bulgu:")
    add_bullet(doc, "Dürüst sınır: Uygulama canlı veri veya tahmin sunmuyor; doğrulanmış tarihsel gözlemleri etkileşimli biçimde açıklıyor.", bold_prefix="Dürüst sınır:")

    # PAGE 2
    add_page_break(doc)
    add_page_header(doc, "01", "Mimari ve veri kaynakları", "Kaynakların görevleri ayrıdır; ortak anahtar timezone-aware UTC timestamp’tir.")
    add_pipeline(doc)
    add_section_heading(doc, "Dört kaynak, dört ayrı sinyal")
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    headers = ["Kaynak", "Sinyal", "Ham çözünürlük", "Birim / alan", "Projede kullanım"]
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        set_cell_shading(c, NAVY)
        set_cell_border(c, WHITE, "8")
        set_cell_margins(c, top=75, bottom=75)
        table_text(c, h, bold=True, color=WHITE, size=7.2, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_table_header(table.rows[0])
    source_rows = [
        ("EnergyZero", "Hollanda saatlik fiyatı", "Saatlik", "€/kWh", "Hedef ölçü; ucuz saat bayrağı"),
        ("ENTSO-E", "Kaynağa göre üretim", "Saatlik", "MW", "Üretim karması ve renewable_share"),
        ("Open-Meteo", "Rüzgâr + güneş", "Saatlik", "m/s, W/m²", "Açıklayıcı değişken ve ML girdisi"),
        ("Wattnet", "Yaşam döngüsü karbon", "15 dakika", "gCO₂e/kWh", "Düşük karbon ölçütü; saatlik ortalama"),
    ]
    for r, row in enumerate(source_rows, start=1):
        cells = table.add_row().cells
        fill = PALE_BLUE if r % 2 else WHITE
        for i, value in enumerate(row):
            set_cell_shading(cells[i], fill)
            set_cell_border(cells[i], LINE, "5")
            set_cell_margins(cells[i], top=70, bottom=70)
            table_text(cells[i], value, bold=(i == 0), size=7.25)

    add_section_heading(doc, "Wattnet karbon çağrısı ve güvenli kimlik doğrulama")
    add_bullet(doc, "GET /v1/footprints; zone=NL, footprint_type=carbon, scope=life-cycle, aggregate=false, use_global=true.")
    add_bullet(doc, "start/end değerleri timezone-aware UTC ISO 8601 olarak gönderiliyor; API yanıtındaki metadata, valid ve zone_status alanları doğrulanıyor.")
    add_bullet(doc, "Token veya geçici cookie yalnızca git-ignore edilen .env dosyasından okunuyor; sırlar kodda veya repoda tutulmuyor.")

    add_section_heading(doc, "Kaynak güvenilirliği nasıl ele alındı?")
    rel = doc.add_table(rows=2, cols=2)
    rel.alignment = WD_TABLE_ALIGNMENT.CENTER
    rel.autofit = True
    rel_data = [
        ("Şema + metadata doğrulaması", "Beklenen sütunlar, birimler, zaman aralığı, çözünürlük ve durum alanları kontrol edildi."),
        ("Kaynaklar arası anlam ayrımı", "Karbon ölçütü Wattnet’tir; yenilenebilir payı açıklayıcıdır ve karbonun yerine kullanılmaz."),
        ("UTC-first yaklaşımı", "DST kaynaklı yerel saat tekrarları birleşim anahtarına sokulmadı; yerel saat yalnızca sunum katmanında üretildi."),
        ("Traceability", "Ham dosyalar ayrı saklandı; temiz tablo ve dashboard export’u aynı dönüşüm zincirinden üretildi."),
    ]
    for idx, (title, text) in enumerate(rel_data):
        cell = rel.cell(idx // 2, idx % 2)
        set_cell_shading(cell, PALE_GREEN if idx % 2 else PALE_BLUE)
        set_cell_border(cell, LINE, "6")
        set_cell_margins(cell, top=85, bottom=85)
        p = cell.paragraphs[0]
        add_run(p, title, bold=True, color=NAVY, size=8)
        p = cell.add_paragraph()
        style_paragraph(p, space_after=0)
        add_run(p, text, size=7.45)

    add_callout(
        doc,
        "Teknik karar",
        "Ham kaynakları doğrudan aynı granülerlikteymiş gibi birleştirmedim. Önce her kaynağı kendi anlamına uygun biçimde temizleyip saatliğe getirdim; sonra inner join ile yalnızca tam eşleşen saatleri analiz ettim.",
        fill=PALE_YELLOW,
        accent=LIME,
    )

    # PAGE 3
    add_page_break(doc)
    add_page_header(doc, "02", "Veri mühendisliği ve MySQL", "Amaç: tekrar çalıştırılabilir, izlenebilir ve güvenli bir ETL zinciri kurmak.")

    add_section_heading(doc, "Uçtan uca dönüşüm sırası")
    steps = [
        "Ingest — HTTP durumunu, JSON şemasını, zorunlu alanları ve değer aralıklarını doğrula; raw CSV’leri ayrı kaydet.",
        "Normalize — timestamp alanlarını timezone-aware UTC’ye çevir; sıralama, numeric cast, duplicate ve negatif değer kontrollerini uygula.",
        "Aggregate — Wattnet 15 dakikalık değerlerini saatlik ortalamaya getir; count=4 olmayan saatleri dışla. Üretim ve hava verisini saatlik hizala.",
        "Merge — fiyat, üretim, hava ve karbonu UTC timestamp üzerinde inner join ile birleştir; eksik kaynaklı saatleri temiz tablodan çıkar.",
        "Engineer — total_generation, renewable_generation, renewable_share ve tarihsel eşik bayraklarını üret.",
        "Load + QA — MySQL şemasını idempotent biçimde oluştur/güncelle; tabloyu yükle, SQL view ve kalite kontrollerini çalıştır.",
    ]
    for s in steps:
        add_numbered(doc, s, size=8.15)

    add_section_heading(doc, "Ana veri sözlüğü")
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, h in enumerate(("Alan", "Tanım", "Kontrol / kullanım")):
        c = table.rows[0].cells[i]
        set_cell_shading(c, NAVY)
        set_cell_border(c, WHITE, "8")
        table_text(c, h, bold=True, color=WHITE, size=7.35, align=WD_ALIGN_PARAGRAPH.CENTER)
    fields = [
        ("timestamp", "UTC saat başlangıcı", "Primary/unique zaman anahtarı"),
        ("electricity_price", "Saatlik enerji fiyatı", "Ucuz = alt çeyrek"),
        ("carbon_intensity_gco2_kwh", "Saatlik yaşam döngüsü yoğunluğu", "Düşük karbon = alt çeyrek"),
        ("renewable_share", "Yenilenebilir üretim / toplam üretim", "0–1 aralığı; açıklayıcı değişken"),
        ("is_cheap_and_low_carbon", "İki alt-çeyrek koşulunun kesişimi", "KPI ve saat örüntüsü"),
        ("wind_speed, solar_radiation", "Saatlik hava değişkenleri", "EDA ve next-hour price model girdisi"),
    ]
    for r, row in enumerate(fields):
        cells = table.add_row().cells
        fill = WHITE if r % 2 else PALE_BLUE
        for i, value in enumerate(row):
            set_cell_shading(cells[i], fill)
            set_cell_border(cells[i], LINE, "5")
            set_cell_margins(cells[i], top=60, bottom=60)
            table_text(cells[i], value, bold=(i == 0), size=7.15)

    add_section_heading(doc, "MySQL ve SQL katmanında neler var?")
    db = doc.add_table(rows=2, cols=2)
    db.alignment = WD_TABLE_ALIGNMENT.CENTER
    db.autofit = True
    db_items = [
        ("Güvenli bağlantı", "SQLAlchemy + PyMySQL; kullanıcı, parola, host ve veritabanı .env’den alınır."),
        ("İdempotent şema", "ensure_hourly_data_schema yeni karbon sütununu yalnızca eksikse ekler; tekrar çalıştırma güvenlidir."),
        ("Analitik view’lar", "vw_analysis_thresholds, vw_hourly_dashboard, vw_daily_summary_utc ve vw_hourly_summary_utc."),
        ("Kalite kontrolleri", "Satır/unique sayısı, duplicate, null, aralık, eşik ve flag tutarlılığı SQL ile doğrulanır."),
    ]
    for idx, (title, text) in enumerate(db_items):
        cell = db.cell(idx // 2, idx % 2)
        set_cell_shading(cell, PALE_GREEN if idx in (1, 2) else PALE_BLUE)
        set_cell_border(cell, LINE, "6")
        set_cell_margins(cell, top=85, bottom=85)
        p = cell.paragraphs[0]
        add_run(p, title, bold=True, color=NAVY, size=8)
        p = cell.add_paragraph()
        add_run(p, text, size=7.4)

    add_callout(
        doc,
        "Önemli implementasyon detayı",
        "Pandas ve SQL eşikleri aynı lineer quantile mantığıyla hesaplanıyor. Böylece notebook, veritabanı view’ı, Tableau export’u ve Streamlit aynı ‘ucuz’ / ‘düşük karbon’ tanımını kullanıyor.",
        fill=PALE_YELLOW,
        accent=BLUE,
    )

    # PAGE 4
    add_page_break(doc)
    add_page_header(doc, "03", "Analiz, hipotez ve makine öğrenmesi", "İstatistiksel bulgu ile ürünün kullandığı tarihsel karar mantığı birbirinden açıkça ayrıldı.")

    add_section_heading(doc, "Hipotez ve tarihsel bulgu")
    add_callout(
        doc,
        "Başlangıç hipotezi",
        "Öğle saatlerinin, güneş üretiminin etkisiyle daha düşük fiyat ve daha düşük karbon yoğunluğunu birlikte sunması bekleniyordu.",
        fill=PALE_BLUE,
        accent=BLUE,
    )
    add_card_grid(
        doc,
        [
            {"value": "12–16", "label": "en iyi tarihsel 4 saat"},
            {"value": "164", "label": "iki koşulu sağlayan saat"},
            {"value": "r=0.495", "label": "fiyat–karbon ilişkisi"},
            {"value": "r=-0.710", "label": "yenilenebilir–karbon ilişkisi"},
        ],
    )
    p = doc.add_paragraph()
    style_paragraph(p, space_before=3, space_after=4)
    add_run(p, "Tanım: ", bold=True, color=NAVY, size=8.2)
    add_run(p, "‘Ucuz’ ve ‘düşük karbon’ ayrı ayrı kendi alt çeyreklerinde (Q1) bulunan saatlerdir; 164 saat iki koşulu aynı anda sağlamıştır. Korelasyon ilişkiyi gösterir, nedensellik kanıtlamaz.", size=8.2)

    add_section_heading(doc, "Daha adil karşılaştırma: aynı günlerde öğle ve akşam")
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, h in enumerate(("Ölçü", "12:00–16:00", "18:00–22:00", "Gözlenen fark")):
        c = table.rows[0].cells[i]
        set_cell_shading(c, NAVY)
        set_cell_border(c, WHITE, "8")
        table_text(c, h, bold=True, color=WHITE, size=7.35, align=WD_ALIGN_PARAGRAPH.CENTER)
    comparisons = [
        ("Ortalama fiyat", "€0.0327/kWh", "€0.1982/kWh", "%83.5 daha düşük"),
        ("Karbon yoğunluğu", "371.7 gCO₂e/kWh", "417.2 gCO₂e/kWh", "%10.9 daha düşük"),
        ("Tarih tutarlılığı", "49/49 günde fiyat düşük", "—", "31/49 günde karbon düşük"),
    ]
    for r, row in enumerate(comparisons):
        cells = table.add_row().cells
        fill = PALE_BLUE if r % 2 == 0 else WHITE
        for i, value in enumerate(row):
            set_cell_shading(cells[i], fill)
            set_cell_border(cells[i], LINE, "5")
            set_cell_margins(cells[i], top=70, bottom=70)
            table_text(cells[i], value, bold=(i in (0, 3)), size=7.3)
    p = doc.add_paragraph()
    style_paragraph(p, space_before=2, space_after=4)
    add_run(p, "49 eksiksiz eşleşmiş gün kullanıldı. Eşleşmiş ortalama farkı için %95 t güven aralığı: fiyat €0.1489–€0.1820/kWh; karbon 15.98–75.13 gCO₂e/kWh.", color=MUTED, size=7.5)

    add_section_heading(doc, "Next-hour price modeli")
    model = doc.add_table(rows=1, cols=2)
    model.alignment = WD_TABLE_ALIGNMENT.CENTER
    model.autofit = False
    model.columns[0].width = Inches(2.0)
    model.columns[1].width = Inches(5.0)
    left, right = model.rows[0].cells
    set_cell_shading(left, NAVY)
    set_cell_border(left, NAVY, "0")
    set_cell_margins(left, top=95, bottom=95)
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "R² 0.943", bold=True, color=LIME, size=15)
    p = left.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "RMSE €0.0189/kWh", bold=True, color=WHITE, size=8)
    p = left.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "210 unseen hours", color=WHITE, size=7.5)
    set_cell_shading(right, PALE_GREEN)
    set_cell_border(right, LINE, "6")
    set_cell_margins(right, top=95, start=125, bottom=95, end=125)
    p = right.paragraphs[0]
    add_run(p, "Değerlendirme tasarımı", bold=True, color=NAVY, size=8.5)
    for text in [
        "Kronolojik hold-out: model geçmiş saatlerde eğitildi, daha sonraki 210 saatte test edildi.",
        "Rastgele split yapılmadı; zaman sızıntısı (leakage) azaltıldı.",
        "R² açıklanan varyanstır; ‘%94.3 doğru tahmin’ anlamına gelmez.",
        "Model proje kapsamında doğrulandı, fakat mevcut Streamlit uygulaması tahmin değil tarihsel veri kullanır.",
    ]:
        p = right.add_paragraph(style="List Bullet")
        style_paragraph(p, space_after=1.5)
        add_run(p, text, size=7.35)

    add_callout(
        doc,
        "Bilimsel ifade",
        "‘Hipotez bu tarihsel örneklemde desteklendi’ demek uygundur. ‘Öğle her zaman en temizdir’ veya ‘uygulama geleceği tahmin ediyor’ demek uygun değildir.",
        fill=PALE_YELLOW,
        accent=RED,
    )

    # PAGE 5
    add_page_break(doc)
    add_page_header(doc, "04", "Ürün, sınırlamalar ve mülakat anlatımı", "Teknik çalışmanın kullanıcıya dönük çıktısı; neyi çözdüğü ve neyi henüz çözmediği.")

    add_section_heading(doc, "Ürün katmanı")
    product = doc.add_table(rows=1, cols=4)
    product.alignment = WD_TABLE_ALIGNMENT.CENTER
    product.autofit = True
    product_items = [
        ("Historical view", "Seçilen geçmiş saat ‘iyi zaman mıydı?’ sorusunu fiyat, karbon ve yenilenebilir payla açıklar."),
        ("Device planner", "Gerçek bir cihazın süre ve tüketimine göre tarihsel maliyet ve CO₂e karşılaştırması yapar."),
        ("Patterns & method", "Hafta günü/saat örüntülerini ve yöntemin sınırlamalarını görünür kılar."),
        ("Tableau", "KPI, saatlik desen, fiyat–karbon ilişkisi ve önerilen pencereyi sunum için özetler."),
    ]
    for idx, (title, text) in enumerate(product_items):
        cell = product.cell(0, idx)
        set_cell_shading(cell, PALE_BLUE if idx % 2 == 0 else PALE_GREEN)
        set_cell_border(cell, LINE, "6")
        set_cell_margins(cell, top=90, bottom=90)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, title, bold=True, color=NAVY, size=7.8)
        p = cell.add_paragraph()
        style_paragraph(p, space_after=0)
        add_run(p, text, size=6.9)

    add_section_heading(doc, "Sınırlamalar ve dürüst teknik yorum")
    limits = [
        "Örneklem geç ilkbahar/yaz dönemidir; kış davranışını ve yıllık mevsimselliği temsil etmez.",
        "Wattnet değeri ortalama yaşam döngüsü karbon yoğunluğudur; marjinal ‘kaçınılan emisyon’ değildir.",
        "Fiyat, saatlik enerji fiyatıdır; ağ/sabit ücretler, sözleşme etkileri ve cihaz verimliliği toplam faturayı değiştirir.",
        "Inner join kaliteyi artırır ancak kapsamı küçültür; eksik saatler sonuçların temsil gücünü etkileyebilir.",
        "164 saat ve 12:00–16:00 sonucu veri setine bağlı tarihsel sınıflandırmadır; canlı tavsiye veya forecast değildir.",
    ]
    for item in limits:
        add_bullet(doc, item, size=7.85)

    add_section_heading(doc, "Bir sonraki teknik adımlar")
    next_table = doc.add_table(rows=1, cols=3)
    next_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    next_table.autofit = True
    next_items = [
        ("Canlılaştır", "Day-ahead fiyat + güncel karbon endpoint’i; cache, retry, rate-limit ve gözlemlenebilirlik ekle."),
        ("Üretimleştir", "ETL orchestration, veri sözleşmeleri, unit/integration test, CI ve model registry kur."),
        ("Genişlet", "Daha uzun/mevsimsel veri, farklı bölgeler, belirsizlik aralıkları ve kullanıcıya özel tarifeler ekle."),
    ]
    for idx, (title, text) in enumerate(next_items):
        cell = next_table.cell(0, idx)
        set_cell_shading(cell, NAVY if idx == 1 else BLUE)
        set_cell_border(cell, WHITE, "8")
        set_cell_margins(cell, top=90, bottom=90)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, title, bold=True, color=LIME, size=8)
        p = cell.add_paragraph()
        add_run(p, text, color=WHITE, size=6.9)

    add_section_heading(doc, "Mülakatta 90 saniyelik anlatım")
    add_callout(
        doc,
        "Önerilen cevap",
        "Bu projede problemimi ‘en ucuz saat’ yerine fiyat ve karbonu birlikte optimize eden tarihsel bir karar sorusu olarak tanımladım. Dört API’den gelen farklı çözünürlükteki verileri UTC’de birleştiren tekrar çalıştırılabilir bir Python ETL’i kurdum. Karbonu 15 dakikadan saatliğe yalnızca tam gözlemlerde agregasyonla getirdim, temiz veriyi MySQL’e idempotent biçimde yükledim ve SQL kalite kontrolleriyle notebook sonuçlarını doğruladım. Analizde alt çeyrek eşikleri, korelasyon ve eşleşmiş gün karşılaştırması kullandım; 12:00–16:00 penceresini bu örneklemde en iyi tarihsel denge olarak buldum. Ayrı olarak leakage’i azaltan kronolojik hold-out ile bir sonraki saat fiyat modelini test ettim. Son olarak bulguyu Streamlit ve Tableau’ya taşıdım; uygulamada canlı tahmin varmış gibi davranmayıp tarihsel kapsamı açıkça gösterdim.",
        fill=PALE_BLUE,
        accent=BLUE,
    )

    add_section_heading(doc, "Beklenen teknik sorulara kısa cevap yönü")
    qa = doc.add_table(rows=2, cols=2)
    qa.alignment = WD_TABLE_ALIGNMENT.CENTER
    qa.autofit = True
    qa_items = [
        ("Neden UTC?", "DST tekrarlarını ve kaynaklar arası saat kaymalarını önlemek; yerel saati yalnızca sunumda üretmek için."),
        ("Neden inner join?", "Analizde her saatte dört sinyalin de bulunmasını garanti etmek için; bedeli daha küçük örneklemdir."),
        ("Neden quantile?", "Ölçekten bağımsız, açıklanabilir ve veri setine uyumlu eşikler üretmek için; mutlak politika eşiği değildir."),
        ("Model neden app’te yok?", "Model doğrulandı ama canlı veri/serving/monitoring zinciri kurulmadığı için üründe tahmin iddiası yapılmadı."),
    ]
    for idx, (q, a) in enumerate(qa_items):
        cell = qa.cell(idx // 2, idx % 2)
        set_cell_shading(cell, WHITE)
        set_cell_border(cell, LINE, "6")
        set_cell_margins(cell, top=70, bottom=70)
        p = cell.paragraphs[0]
        add_run(p, q, bold=True, color=NAVY, size=7.7)
        p = cell.add_paragraph()
        add_run(p, a, size=7.0)

    return doc


if __name__ == "__main__":
    OUT.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUT)
    print(OUT)
